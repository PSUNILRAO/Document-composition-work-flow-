"""
template_studio.py
───────────────────
Target-state "Template Studio" capability.

Responsibilities
────────────────
1. **Placeholder extraction** — given an uploaded DOCX template, parse the
   Jinja2 expressions it contains and return a structured description of:
     * scalar placeholders     — e.g. ``{{ customer_name }}``
     * repeating sections      — e.g. ``{% for t in transactions %}…{% endfor %}``
       with the inner field names referenced against the loop variable.
2. **Binding manifest** — per ``doc_type`` JSON file stored alongside the
   uploaded template. Maps placeholder names in the DOCX to the actual field
   names present in the recipient's data (so business users can drop a DOCX
   authored with ``{{ customer_name }}`` on top of an Excel whose column is
   ``NAME``).
3. **Binding application** — at render time, enrich the record context so every
   placeholder resolves. Scalar bindings are a simple copy; repeating bindings
   re-shape the source list-of-dicts so the inner field names match what the
   DOCX references.

The module is intentionally pure — no Flask/UI concerns live here.
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

from docx import Document
from jinja2 import Environment, meta, nodes

from docx_renderer import UPLOAD_TEMPLATES_DIR, uploaded_template_path

log = logging.getLogger(__name__)

# Bindings manifests live next to the uploaded DOCX, one JSON per doc_type.
BINDINGS_DIR = UPLOAD_TEMPLATES_DIR
_BINDINGS_DIR_RESOLVED = BINDINGS_DIR.resolve()

# Manifest schema version. Bumped when the on-disk format changes in a
# backwards-incompatible way.
MANIFEST_VERSION = 1

# Jinja environment used for AST parsing only. We never execute templates here.
_parse_env = Environment()


# ─────────────────────────────────────────────────────────────────────────────
# Placeholder extraction
# ─────────────────────────────────────────────────────────────────────────────
def _collect_template_source(docx_path: Path) -> str:
    """Return a Jinja-parseable source string assembled from every paragraph
    and table cell in the DOCX. Runs are merged into one string per paragraph,
    mirroring what :mod:`docx_renderer` does at render time so we see the same
    tokens Word users author."""

    doc = Document(str(docx_path))
    parts: list[str] = []

    def _walk_paragraphs(paragraphs):
        for para in paragraphs:
            text = "".join(r.text or "" for r in para.runs)
            if text.strip():
                parts.append(text)

    def _walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _walk_paragraphs(cell.paragraphs)
                    _walk_tables(cell.tables)

    _walk_paragraphs(doc.paragraphs)
    _walk_tables(doc.tables)

    return "\n".join(parts)


def _inner_field_names(body: list, loop_var: str) -> list[str]:
    """Walk a Jinja AST sub-tree and return every attribute name accessed
    against ``loop_var`` (e.g. inside ``{% for t in trades %}{{ t.date }}``
    returns ``["date"]``).

    We only surface a single level of attribute access, which is what authors
    writing DOCX templates realistically use for tabular data. Nested attribute
    access (``t.contact.email``) is returned as its first-level name.
    """

    seen: list[str] = []

    def _walk(node: Any) -> None:
        if isinstance(node, nodes.Getattr):
            target = node.node
            if isinstance(target, nodes.Name) and target.name == loop_var:
                if node.attr not in seen:
                    seen.append(node.attr)
        # Recurse into child nodes
        if hasattr(node, "iter_child_nodes"):
            for child in node.iter_child_nodes():
                _walk(child)

    for stmt in body:
        _walk(stmt)
    return seen


def extract_placeholders(doc_type: str, docx_path: Path | None = None) -> dict[str, Any]:
    """Return a structured description of every placeholder in a DOCX template.

    By default looks up the currently-published DOCX for ``doc_type``
    (``uploaded_template_path(doc_type)``). Pass ``docx_path`` explicitly to
    extract placeholders from an arbitrary version's DOCX (used by the Studio
    version picker).

    Return shape::

        {
          "template_name": "bank_statement.docx",
          "scalar_placeholders": ["customer_name", "closing_balance", ...],
          "repeating_sections": [
            {
              "iter_source":  "transactions",   # the variable iterated over
              "loop_var":     "t",              # name used inside the loop
              "inner_fields": ["date", "description", "amount"]
            }
          ],
          "parse_error": None,   # or a string if the DOCX is not parseable
        }

    Raises ``FileNotFoundError`` if no DOCX is present at the resolved path.
    """

    path = docx_path if docx_path is not None else uploaded_template_path(doc_type)
    if not path.is_file():
        raise FileNotFoundError(f"No DOCX template found at {path}.")

    result: dict[str, Any] = {
        "template_name": path.name,
        "scalar_placeholders": [],
        "repeating_sections": [],
        "parse_error": None,
    }

    try:
        source = _collect_template_source(path)
    except Exception as exc:  # noqa: BLE001 — report any DOCX-level failure
        log.warning("Failed to read DOCX %s: %s", path, exc)
        result["parse_error"] = f"Could not read DOCX: {exc}"
        return result

    try:
        ast = _parse_env.parse(source)
    except Exception as exc:  # noqa: BLE001 — Jinja surfaces many error types
        log.warning("Jinja parse failed for %s: %s", path, exc)
        # Fall back to a regex sweep so the author still sees *something*.
        result["parse_error"] = f"Jinja parse error: {exc}"
        result["scalar_placeholders"] = sorted(_regex_scalar_names(source))
        return result

    # Repeating sections: every {% for X in Y %} block in the AST.
    for_nodes: list[nodes.For] = list(ast.find_all(nodes.For))
    vars_inside_loops: set[str] = set()
    for fn in for_nodes:
        iter_name = _name_of(fn.iter)
        loop_var = _name_of(fn.target) or ""
        inner = _inner_field_names(fn.body, loop_var) if loop_var else []
        if iter_name:
            result["repeating_sections"].append({
                "iter_source":  iter_name,
                "loop_var":     loop_var,
                "inner_fields": inner,
            })
        # Track identifiers that live inside the loop so we don't double-count
        # them as scalar placeholders.
        for child in fn.iter_child_nodes():
            for n in child.find_all(nodes.Name):
                vars_inside_loops.add(n.name)
        if loop_var:
            vars_inside_loops.add(loop_var)

    # Scalar placeholders: every top-level undeclared variable that isn't an
    # iter-source or loop variable.
    all_undeclared = meta.find_undeclared_variables(ast)
    iter_sources = {rs["iter_source"] for rs in result["repeating_sections"]}
    scalars = sorted(
        v for v in all_undeclared
        if v not in iter_sources and v not in vars_inside_loops
    )
    result["scalar_placeholders"] = scalars
    return result


def _name_of(node: Any) -> str | None:
    """Return the dotted-name string for a Jinja AST node, or ``None``."""
    if isinstance(node, nodes.Name):
        return node.name
    if isinstance(node, nodes.Getattr):
        base = _name_of(node.node)
        return f"{base}.{node.attr}" if base else None
    return None


_SCALAR_RE = re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*)")


def _regex_scalar_names(source: str) -> set[str]:
    """Best-effort scalar placeholder extraction when Jinja parsing fails."""
    return set(_SCALAR_RE.findall(source))


# ─────────────────────────────────────────────────────────────────────────────
# Binding manifest storage
# ─────────────────────────────────────────────────────────────────────────────
def _safe_bindings_path(doc_type: str) -> Path | None:
    """Resolve the bindings JSON path for ``doc_type``, refusing traversal."""
    if not doc_type or not isinstance(doc_type, str):
        return None
    if doc_type in (".", "..") or any(c in doc_type for c in ("/", "\\", "\x00")):
        return None
    candidate = (BINDINGS_DIR / f"{doc_type}.bindings.json").resolve()
    try:
        candidate.relative_to(_BINDINGS_DIR_RESOLVED)
    except ValueError:
        return None
    return candidate


def _empty_manifest() -> dict[str, Any]:
    return {
        "version":    MANIFEST_VERSION,
        "scalars":    {},   # placeholder_name -> source_field_name
        "repeating":  {},   # iter_source -> { "source": <data list>,
                            #                  "field_map": { inner: source_field } }
    }


def load_bindings(doc_type: str) -> dict[str, Any]:
    """Return the saved binding manifest for ``doc_type`` (empty if none)."""
    path = _safe_bindings_path(doc_type)
    if path is None or not path.is_file():
        return _empty_manifest()
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except (OSError, json.JSONDecodeError) as exc:
        log.warning("Bad bindings manifest at %s: %s", path, exc)
        return _empty_manifest()

    # Normalise to the current shape so the rest of the code can assume it.
    manifest = _empty_manifest()
    if isinstance(data.get("scalars"), dict):
        for k, v in data["scalars"].items():
            if isinstance(k, str) and isinstance(v, str):
                manifest["scalars"][k] = v
    if isinstance(data.get("repeating"), dict):
        for k, v in data["repeating"].items():
            if not (isinstance(k, str) and isinstance(v, dict)):
                continue
            source = v.get("source")
            field_map = v.get("field_map") or {}
            if not isinstance(source, str) or not isinstance(field_map, dict):
                continue
            manifest["repeating"][k] = {
                "source":    source,
                "field_map": {ik: iv for ik, iv in field_map.items()
                              if isinstance(ik, str) and isinstance(iv, str)},
            }
    return manifest


def normalise_manifest(manifest: dict[str, Any]) -> dict[str, Any]:
    """Strip junk/empty entries and enforce types on a bindings manifest."""
    if not isinstance(manifest, dict):
        raise ValueError("manifest must be a dict")

    normalised = _empty_manifest()
    scalars = manifest.get("scalars") or {}
    if isinstance(scalars, dict):
        for k, v in scalars.items():
            if isinstance(k, str) and isinstance(v, str) and v:
                normalised["scalars"][k] = v
    repeating = manifest.get("repeating") or {}
    if isinstance(repeating, dict):
        for k, v in repeating.items():
            if not (isinstance(k, str) and isinstance(v, dict)):
                continue
            source = v.get("source")
            if not isinstance(source, str) or not source:
                continue
            field_map = v.get("field_map") or {}
            cleaned_map: dict[str, str] = {}
            if isinstance(field_map, dict):
                for ik, iv in field_map.items():
                    if isinstance(ik, str) and isinstance(iv, str) and iv:
                        cleaned_map[ik] = iv
            normalised["repeating"][k] = {
                "source":    source,
                "field_map": cleaned_map,
            }
    return normalised


def save_bindings(doc_type: str, manifest: dict[str, Any]) -> dict[str, Any]:
    """Persist a binding manifest. Returns the normalised manifest as saved."""
    path = _safe_bindings_path(doc_type)
    if path is None:
        raise ValueError(f"Invalid doc_type for bindings path: {doc_type!r}")

    normalised = normalise_manifest(manifest)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(normalised, f, indent=2, sort_keys=True)
    return normalised


def remove_bindings(doc_type: str) -> bool:
    """Delete the saved manifest for ``doc_type`` if it exists."""
    path = _safe_bindings_path(doc_type)
    if path is not None and path.is_file():
        path.unlink()
        return True
    return False


def bindings_exist(doc_type: str) -> bool:
    path = _safe_bindings_path(doc_type)
    return path is not None and path.is_file()


# ─────────────────────────────────────────────────────────────────────────────
# Binding application at render time
# ─────────────────────────────────────────────────────────────────────────────
def apply_bindings(context: dict[str, Any],
                   manifest: dict[str, Any]) -> dict[str, Any]:
    """Return a new context with placeholder names populated per the manifest.

    The original ``context`` keys are preserved (so existing 1:1 templates keep
    working); binding-derived keys are *added*. Where a placeholder name
    collides with an existing key, the binding wins — that's the whole point of
    the mapping.
    """

    if not manifest or (not manifest.get("scalars")
                        and not manifest.get("repeating")):
        return context

    enriched = dict(context)

    # Scalar bindings: placeholder ← source_field
    for placeholder, source_field in (manifest.get("scalars") or {}).items():
        if not placeholder or not source_field:
            continue
        if source_field in enriched:
            enriched[placeholder] = enriched[source_field]

    # Repeating bindings: reshape each source list into a list of dicts whose
    # keys match what the DOCX loop body references.
    for iter_name, spec in (manifest.get("repeating") or {}).items():
        if not iter_name or not isinstance(spec, dict):
            continue
        source = spec.get("source")
        field_map = spec.get("field_map") or {}
        if not source or source not in enriched:
            continue
        source_rows = enriched.get(source) or []
        if not isinstance(source_rows, list):
            continue
        reshaped: list[dict[str, Any]] = []
        for row in source_rows:
            if not isinstance(row, dict):
                # Non-dict rows can still be surfaced as a single "value" field
                # so ``{{ item }}`` style loops keep working.
                reshaped.append({"value": row})
                continue
            new_row = dict(row)  # carry everything through by default
            for inner_placeholder, inner_source in field_map.items():
                if inner_source in row:
                    new_row[inner_placeholder] = row[inner_source]
            reshaped.append(new_row)
        enriched[iter_name] = reshaped

    return enriched
