"""
docx_renderer.py
─────────────────
Render a user-uploaded .docx template to PDF bytes, merging in record data.

How it works
────────────
1. Load the DOCX with python-docx.
2. Substitute Jinja2 placeholders — `{{field}}`, `{% if … %}` etc. — in every
   paragraph and table cell, using the same enriched record context (record +
   rule-engine output) that the HTML renderer uses.
3. Walk the document body in document order and build a ReportLab PDF that
   preserves paragraphs and tables.

Dependencies
────────────
Only `python-docx`, `jinja2`, and `reportlab` — all already listed in
requirements.txt. No LibreOffice or external binaries required.
"""

from __future__ import annotations

import io
import logging
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as _DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as _DocxTable
from docx.text.paragraph import Paragraph as _DocxParagraph
from jinja2 import Environment, Undefined

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

log = logging.getLogger(__name__)

# Uploaded DOCX templates live here (per doc_type). The engine picks the file
# up automatically when generating; UI manages add/remove.
UPLOAD_TEMPLATES_DIR = Path(__file__).parent / "uploads" / "templates"
UPLOAD_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


# ── Jinja2 environment ────────────────────────────────────────────────────────
class _SilentUndefined(Undefined):
    """Undefined placeholders render as empty string instead of raising."""

    def __str__(self) -> str:  # noqa: D401
        return ""


_jinja = Environment(
    undefined=_SilentUndefined,
    autoescape=False,            # DOCX runs are plain text, not HTML
    keep_trailing_newline=True,
)

# Reuse the same formatting filters the HTML renderer exposes, so DOCX authors
# can write `{{ closing_balance | currency }}` just like in the built-in HTML
# templates. Imported lazily to avoid a circular import at module load time.
def _install_shared_filters() -> None:
    from renderer import (_fmt_currency, _fmt_number, _fmt_percent,
                          _row_style)  # noqa: PLC0415
    _jinja.filters["currency"]  = _fmt_currency
    _jinja.filters["percent"]   = _fmt_percent
    _jinja.filters["number"]    = _fmt_number
    _jinja.filters["row_style"] = _row_style


_install_shared_filters()


# ── Template-path helpers ─────────────────────────────────────────────────────
_UPLOAD_TEMPLATES_DIR_RESOLVED = UPLOAD_TEMPLATES_DIR.resolve()


def _safe_template_path(doc_type: str) -> Path | None:
    """Resolve the DOCX template path for ``doc_type``.

    Returns ``None`` if ``doc_type`` is empty, contains path separators, or
    would resolve outside of :data:`UPLOAD_TEMPLATES_DIR`. This is
    defense-in-depth: callers are expected to have already validated
    ``doc_type`` against an allow-list, but hardening here ensures these
    helpers cannot be used to reach arbitrary filesystem paths if any caller
    forgets to validate.
    """
    if not doc_type or not isinstance(doc_type, str):
        return None
    # Reject anything that isn't a plain identifier-ish name. This is stricter
    # than the allow-list check (which is exact-match against DOC_LABELS) but
    # guarantees no traversal or separator characters can reach the filesystem.
    if doc_type in (".", "..") or any(c in doc_type for c in ("/", "\\", "\x00")):
        return None
    candidate = (UPLOAD_TEMPLATES_DIR / f"{doc_type}.docx").resolve()
    try:
        candidate.relative_to(_UPLOAD_TEMPLATES_DIR_RESOLVED)
    except ValueError:
        return None
    return candidate


def uploaded_template_path(doc_type: str) -> Path:
    """Where we store the uploaded DOCX for a given doc_type.

    Raises ``ValueError`` if ``doc_type`` would resolve outside of
    :data:`UPLOAD_TEMPLATES_DIR`. Callers that want a soft-failing check should
    use :func:`has_uploaded_template` (which returns ``False`` instead).
    """
    safe = _safe_template_path(doc_type)
    if safe is None:
        raise ValueError(f"Invalid doc_type for template path: {doc_type!r}")
    return safe


def has_uploaded_template(doc_type: str) -> bool:
    safe = _safe_template_path(doc_type)
    return safe is not None and safe.is_file()


def remove_uploaded_template(doc_type: str) -> bool:
    safe = _safe_template_path(doc_type)
    if safe is not None and safe.is_file():
        safe.unlink()
        return True
    return False


# ── Placeholder substitution ──────────────────────────────────────────────────
def _render_text(text: str, context: dict) -> str:
    if "{{" not in text and "{%" not in text:
        return text
    try:
        return _jinja.from_string(text).render(**context)
    except Exception as exc:  # noqa: BLE001 — template authors may typo
        log.warning("docx template render error (%s): %r", exc, text[:80])
        return text


def _merge_paragraph(para: _DocxParagraph, context: dict) -> None:
    """
    Replace {{ }} / {% %} tokens inside a paragraph.

    Word often splits tokens across multiple `run` elements (because of
    formatting boundaries), so we merge the runs into one string, render it,
    and write the result into the first run while clearing the others.
    """
    runs = list(para.runs)
    if not runs:
        return
    full_text = "".join(r.text or "" for r in runs)
    if "{{" not in full_text and "{%" not in full_text:
        return
    rendered = _render_text(full_text, context)
    if rendered == full_text:
        return
    runs[0].text = rendered
    for r in runs[1:]:
        r.text = ""


def _merge_table(table: _DocxTable, context: dict) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _merge_paragraph(para, context)
            for inner in cell.tables:  # nested tables
                _merge_table(inner, context)


# ── Multi-paragraph block expansion ───────────────────────────────────────────
_BLOCK_START_RE = re.compile(r"{%\s*(for|if)\b")
_BLOCK_END_RE   = re.compile(r"{%\s*end(for|if)\s*%}")


def _para_text(p: _DocxParagraph) -> str:
    return "".join(r.text or "" for r in p.runs)


def _set_paragraph_text(p: _DocxParagraph, text: str) -> None:
    runs = list(p.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _expand_multi_paragraph_blocks(doc: _DocxDocument, context: dict) -> None:
    """Render ``{% for %}…{% endfor %}`` (or ``if``) blocks that span multiple
    paragraphs.

    python-docx represents each visual line as a separate ``<w:p>`` element, so
    a DOCX authored with the loop body on its own line cannot be rendered
    paragraph-by-paragraph (the opening ``{% for %}`` tag on one paragraph and
    the matching ``{% endfor %}`` on another are each invalid Jinja in
    isolation). This pass detects such blocks, joins their text with newlines,
    renders the whole span as one Jinja template, and splits the result back
    into one paragraph per output line — cloning the formatting of the first
    paragraph in the block so downstream PDF rendering preserves the author's
    styling.
    """
    body = doc.element.body

    def top_level_paragraphs() -> list[_DocxParagraph]:
        return [_DocxParagraph(p, doc) for p in body.findall(qn("w:p"))]

    # Each pass handles a single outermost block; re-scan after mutation to
    # stay in sync with newly-created paragraphs.
    safety = 0
    while safety < 1000:
        safety += 1
        paragraphs = top_level_paragraphs()
        target: tuple[int, int] | None = None
        for i, p in enumerate(paragraphs):
            text = _para_text(p)
            if not _BLOCK_START_RE.search(text):
                continue
            if _BLOCK_END_RE.search(text):
                continue  # fully self-contained in one paragraph
            depth = 1
            j = i + 1
            while j < len(paragraphs) and depth > 0:
                tj = _para_text(paragraphs[j])
                has_start = bool(_BLOCK_START_RE.search(tj))
                has_end = bool(_BLOCK_END_RE.search(tj))
                if has_start and has_end:
                    # Self-contained inner block on a single paragraph
                    # (e.g. ``{% if x.active %}…{% endif %}``) — neither opens
                    # nor closes the outer span, so leave depth untouched.
                    pass
                elif has_start:
                    depth += 1
                elif has_end:
                    depth -= 1
                    if depth == 0:
                        break
                j += 1
            if j < len(paragraphs) and depth == 0:
                target = (i, j)
                break
        if target is None:
            return

        i, j = target
        block = paragraphs[i:j + 1]
        template_src = "\n".join(_para_text(bp) for bp in block)
        rendered = _render_text(template_src, context)
        # Collapse whitespace-only lines left behind by the for/endfor tags.
        lines = [ln for ln in rendered.split("\n") if ln.strip() != ""]
        if not lines:
            lines = [""]

        first = block[0]
        _set_paragraph_text(first, lines[0])
        anchor = first._element
        for extra in lines[1:]:
            clone = deepcopy(first._element)
            anchor.addnext(clone)
            anchor = clone
            _set_paragraph_text(_DocxParagraph(clone, doc), extra)
        for bp in block[1:]:
            bp._element.getparent().remove(bp._element)

    log.warning("multi-paragraph block expansion hit safety cap")


def merge_docx(docx_path: str | Path, context: dict) -> _DocxDocument:
    """Load a DOCX template and substitute Jinja2 placeholders in-place."""
    doc = Document(str(docx_path))
    _expand_multi_paragraph_blocks(doc, context)
    for para in doc.paragraphs:
        _merge_paragraph(para, context)
    for table in doc.tables:
        _merge_table(table, context)
    return doc


# ── PDF rendering ─────────────────────────────────────────────────────────────
def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Heading 1": ParagraphStyle(
            "H1", parent=base["Title"],
            fontSize=18, textColor=colors.HexColor("#1E3A5F"),
            fontName="Helvetica-Bold", spaceAfter=6, alignment=TA_LEFT,
        ),
        "Heading 2": ParagraphStyle(
            "H2", parent=base["Heading2"],
            fontSize=13, textColor=colors.HexColor("#1E3A5F"),
            fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4,
        ),
        "Heading 3": ParagraphStyle(
            "H3", parent=base["Heading3"],
            fontSize=11, textColor=colors.HexColor("#1E3A5F"),
            fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=3,
        ),
        "Normal": ParagraphStyle(
            "Body", parent=base["Normal"],
            fontSize=10, leading=14, fontName="Helvetica",
            textColor=colors.HexColor("#1F2937"), spaceAfter=4,
        ),
        "TableHeader": ParagraphStyle(
            "TH", parent=base["Normal"],
            fontSize=9, fontName="Helvetica-Bold", textColor=colors.white,
        ),
        "TableCell": ParagraphStyle(
            "TD", parent=base["Normal"],
            fontSize=9, fontName="Helvetica",
            textColor=colors.HexColor("#1F2937"),
        ),
    }


def _xml_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
    )


def _paragraph_story(para: _DocxParagraph, styles: dict) -> list:
    text = (para.text or "").strip()
    if not text:
        return [Spacer(1, 4)]
    style_name = para.style.name if para.style else "Normal"
    style = styles.get(style_name) or styles["Normal"]
    # Render a few basic inline styles (bold/italic) from the DOCX runs.
    pieces: list[str] = []
    for run in para.runs:
        t = _xml_escape(run.text or "")
        if not t:
            continue
        if run.bold:
            t = f"<b>{t}</b>"
        if run.italic:
            t = f"<i>{t}</i>"
        if run.underline:
            t = f"<u>{t}</u>"
        pieces.append(t)
    rendered = "".join(pieces) or _xml_escape(text)
    return [Paragraph(rendered, style)]


def _table_story(table: _DocxTable, styles: dict, usable_w: float) -> list:
    tdata = []
    for r_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            txt = "\n".join(p.text for p in cell.paragraphs).strip()
            style = styles["TableHeader"] if r_idx == 0 else styles["TableCell"]
            cells.append(Paragraph(_xml_escape(txt) or "&nbsp;", style))
        tdata.append(cells)
    if not tdata:
        return []
    col_n = max(len(r) for r in tdata)
    col_w = usable_w / col_n
    # Pad short rows so ReportLab doesn't choke.
    for r in tdata:
        while len(r) < col_n:
            r.append(Paragraph("", styles["TableCell"]))
    t = Table(tdata, colWidths=[col_w] * col_n, hAlign="LEFT", repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0),  colors.HexColor("#1E3A5F")),
        ("TEXTCOLOR",      (0, 0), (-1, 0),  colors.white),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F0F4FF")]),
        ("GRID",           (0, 0), (-1, -1), 0.3, colors.HexColor("#E5E7EB")),
        ("TOPPADDING",     (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",  (0, 0), (-1, -1), 5),
        ("LEFTPADDING",    (0, 0), (-1, -1), 7),
        ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return [Spacer(1, 4), t, Spacer(1, 8)]


def _iter_body_blocks(doc: _DocxDocument):
    """Yield ('paragraph', p) and ('table', t) in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", _DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "table", _DocxTable(child, doc)


def docx_to_pdf(doc: _DocxDocument) -> bytes:
    """Render a merged python-docx Document to PDF bytes via ReportLab."""
    styles = _pdf_styles()
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
    )
    usable_w = A4[0] - 40 * mm
    story: list = []
    for kind, block in _iter_body_blocks(doc):
        if kind == "paragraph":
            story.extend(_paragraph_story(block, styles))
        else:
            story.extend(_table_story(block, styles, usable_w))
    if not story:
        story.append(Paragraph(
            "DOCX template rendered with no visible content.",
            styles["Normal"],
        ))
    pdf.build(story)
    return buf.getvalue()


# ── Public entry point ────────────────────────────────────────────────────────
def render_docx_pdf(docx_path: str | Path, context: dict) -> bytes:
    """Top-level: DOCX template path + merge context → PDF bytes."""
    merged = merge_docx(docx_path, context)
    return docx_to_pdf(merged)
